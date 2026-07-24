from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("人员组成及分工（已填写）.docx")

ROWS = [
    ("1", "项目负责人（待明确）", "车能处", "负责项目总体统筹、进度管控、跨部门协调和产业链沟通，组织阶段成果审查与验收。"),
    ("2", "标准与体系负责人（待明确）", "车能处", "负责BMS通信标准体系搭建、关键节点把控，协调标准框架、专家评审及发布推广工作。"),
    ("3", "组织协调负责人（待明确）", "车能处", "负责委内协调、调研组织、试点单位联络及项目过程资料归集，协助项目进度把控。"),
    ("4", "文本与质量负责人（待明确）", "车能处", "负责项目全流程文本资料、技术规范和验收材料的审核把关，确保成果符合管理及标准化要求。"),
    ("5", "技术负责人（待明确）", "全来电", "负责总体技术路线、BMS通信架构和关键技术决策；组织UART/CAN/BLE接口、数据链路层和应用层协议设计。"),
    ("6", "协议与测试工程师（待明确）", "全来电", "负责核心信号模型、帧格式、CRC校验、流控及异常处理设计；开发协议一致性测试工具，开展多品牌联调与测试报告编制。"),
    ("7", "车辆企业代表（待明确）", "车辆企业", "提供车辆侧控制器与直流充电接口需求、通信数据及样车验证条件，参与接口适配、联调测试和试点验证。"),
    ("8", "电池企业代表（待明确）", "电池企业", "提供BMS核心信号、身份认证和安全保护需求，参与电池侧协议适配、充电参数协商及互通验证。"),
    ("9", "换电/充电场景代表（待明确）", "全来电", "负责换电柜、充电器及IoT终端通信需求分析，参与换电互通、动态充电策略和异常停充场景验证。"),
    ("10", "科研机构专家（待明确）", "科研机构", "负责技术方案咨询、车池桩数据协同验证、测试结果评估及标准规范专家评审。"),
]


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def write_cell(cell, text, *, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    title_run = title.add_run("人员组成及分工")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("电动自行车直流充电路径研究（科技类课题）")
    subtitle_run.font.name = "Microsoft YaHei"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(89, 102, 115)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [800, 2300, 1500, 4760]
    set_table_geometry(table, widths)
    headers = ["序号", "姓名/岗位", "单位", "分工"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8EEF5")
        write_cell(cell, header, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)

    for index, name, unit, duty in ROWS:
        cells = table.add_row().cells
        write_cell(cells[0], index, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(cells[1], name, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(cells[3], duty)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("说明：课题研究文件未提供参研人员真实姓名，本文以岗位名称“待明确”代填；确定人员后请在“姓名/岗位”列替换为真实姓名。")
    note_run.font.name = "Microsoft YaHei"
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(89, 102, 115)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
